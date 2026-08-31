import os
import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing it...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    import docx

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_comprehensive_document():
    print("Generating comprehensive Project_Overview_Threat_Detection.docx...")
    doc = docx.Document()
    
    # -------------------------------------------------------------
    # DOCUMENT GEOMETRY
    # -------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # -------------------------------------------------------------
    # STYLES DEFINITION
    # -------------------------------------------------------------
    # Normal body text
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate gray
    
    # Title Page Paragraphs
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("TECHNICAL OVERVIEW & HANDBOOK: SOFTWARE SUPPLY-CHAIN THREAT DETECTION GATEKEEPER")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a) # Slate 900
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(8)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("An AI-Driven Chronological Multimodal Framework (GNN + LSTM)")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    subtitle_p.paragraph_format.space_after = Pt(24)
    
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.add_run("━" * 45).font.color.rgb = RGBColor(0xcb, 0xd5, 0xe1)
    p_div.paragraph_format.space_after = Pt(24)
    
    # Helpers for headings & paragraphs
    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a) # Slate 900
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a) # Blue 900
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p

    def add_body_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
        run = p.add_run(text)
        run.italic = italic
        return p

    def add_bullet_p(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
        p.add_run(text)
        return p

    def add_formula_p(formula_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(formula_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        # Add light gray background border shading visually by borders (handled in word default spacing)
        return p

    # -------------------------------------------------------------
    # SECTION 1: ARCHITECTURE AND FILE LAYOUT
    # -------------------------------------------------------------
    add_heading_1("1. System Architecture & Path Resolutions")
    add_body_p(
        "To ensure stability under different run configurations, the system uses absolute directory anchoring. "
        "The following table maps the core components, their file paths, and their specific technical functions in the codebase:"
    )
    
    # Codebase Paths Table
    paths_headers = ["Component", "File Path", "Operational Description"]
    paths_data = [
        ["Main Controller", "main.py", "Acts as the central HTTP server handling REST endpoints. Initiates threat prediction pipelines, seeds sqlite tables, and triggers host installations."],
        ["Dependency Scanner", "dependency_analysis/dependency_scanner.py", "Performs transit dependency resolution, queries OSV.dev APIs, checks NVD Excel databases, and logs ASCII risk trees to the console."],
        ["GNN Classifier", "ml_engine/hgat_model.py", "Implements the Heterogeneous Graph Attention Network (HGAT) architecture in PyTorch."],
        ["LSTM Model", "ml_engine/temporal_model.py", "Implements the Temporal LSTM classifier designed to identify sequential release anomalies."],
        ["Training Engine", "ml_engine/trainer.py", "Handles separate Training split (<=2022) and Validation split (2023-2024) performance evaluations."],
        ["Static Analyzer", "static_analysis/analyzer.py", "Checks files via AST analysis, measures obfuscation entropy, and executes Bandit/Semgrep CLI checks on quarantined sources."],
        ["Dashboard template", "visualization/dashboard.py", "Dynamically generates the glassmorphic Security Dashboard HTML file, parsing JSON inputs and bypassing curly-brace formatting conflicts."]
    ]
    
    t_paths = doc.add_table(rows=1, cols=3)
    t_paths.style = 'Light Shading Accent 1'
    hdr_cells = t_paths.rows[0].cells
    for i, h in enumerate(paths_headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F2937")
        for r in hdr_cells[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            
    for row_idx, r_data in enumerate(paths_data):
        row_cells = t_paths.add_row().cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(r_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_col)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 2: ENDPOINT ROUTING AND REST APIS
    # -------------------------------------------------------------
    add_heading_1("2. REST API Specifications & Routing")
    add_body_p(
        "Communication between the glassmorphism frontend and Python backend is mediated via internal REST endpoints. "
        "The following table details the HTTP methods, URI paths, and request-response parameters:"
    )
    
    api_headers = ["Method", "Endpoint Path", "Request / Input", "Response / Output"]
    api_data = [
        ["POST", "/scan_requirements", "Plain-text requirements file body uploaded in dashboard.", "JSON redirect containing scan statuses, GNN risks, and dashboard.html locations."],
        ["GET", "/list_outputs", "Triggered by 'Explore Output Artifacts' button.", "JSON list of whitelisted scan files (e.g. SBOM, Semgrep report, PNG plots) in outputs/ folder."],
        ["GET", "/view_output?file={name}", "Passes whitelisted filename query parameter.", "Streams file contents back. Dynamically sets Content-Type (text/plain or image/png)."],
        ["POST", "/install", "No body parameters. Installs requirements fromoutputs/last_requirements.txt.", "Streams live pip installation output logs. Returns 200 Success or 500 Stacktrace."],
        ["POST", "/abort", "Triggered when user clicks 'Abort' or 'Close Viewer'.", "Cleans staging directories and returns success string."]
    ]
    
    t_api = doc.add_table(rows=1, cols=4)
    t_api.style = 'Light Shading Accent 1'
    hdr_api = t_api.rows[0].cells
    for i, h in enumerate(api_headers):
        hdr_api[i].text = h
        set_cell_background(hdr_api[i], "1F2937")
        for r in hdr_api[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            
    for row_idx, r_data in enumerate(api_data):
        row_cells = t_api.add_row().cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(r_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_col)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 3: MATHEMATICAL FORMULATIONS
    # -------------------------------------------------------------
    add_heading_1("3. Core Security Modalities & Mathematical Formulations")
    add_body_p(
        "Each of the five analysis modalities uses specific mathematical formulas or computational algorithms to model risk parameters."
    )
    
    add_heading_2("3.1 Source Code Obfuscation (Shannon Entropy)")
    add_body_p(
        "During pre-install scanning, we calculate the Shannon Entropy of code literals. High character randomness indicates obfuscated scripts, "
        "encrypted shellcode payloads, or packed malicious code blocks. The Shannon Entropy H(X) is defined as:"
    )
    add_formula_p("H(X) = - sum_{i=1}^{n} P(x_i) * log_2 P(x_i)")
    add_body_p(
        "Where P(x_i) is the probability of occurrence of character x_i in the string sequence X, and n is the size of the unique character vocabulary. "
        "Standard Python source files have a Shannon Entropy score of 3.5 to 4.8. Obfuscated base64 payload sequences typically exceed 5.8."
    )
    
    add_heading_2("3.2 Heterogeneous Graph Attention Network (HGAT)")
    add_body_p(
        "To model dependency risk propagation across transitive dependency trees, our HGAT GNN projects packages and CVE relationships "
        "into a graph structure. The attention coefficients alpha_ij measure how strongly parent package i is affected by dependency node j:"
    )
    add_formula_p("alpha_ij = exp( LeakyReLU( a^T * [ W * h_i || W * h_j ] ) ) / ( sum_{k in N_i} exp( LeakyReLU( a^T * [ W * h_i || W * h_k ] ) ) )")
    add_body_p(
        "Where h_i and h_j are feature representation vectors of nodes i and j, W is the shared parameter projection matrix, a is the weight vector "
        "of the attention mechanism, || denotes the concatenation operator, and N_i is the neighbourhood set of node i. This ensures "
        "that vulnerabilities located deep in transitive dependencies (e.g. urllib3) propagate risk attention weights up to application levels."
    )
    
    add_heading_2("3.3 LSTM Sequence Anomaly Modeling")
    add_body_p(
        "A package account takeover is modeled as a temporal release interval anomaly. The LSTM inspects chronological release delay sequences "
        "delta_t = t_k - t_{k-1}. The LSTM cell updates are governed by standard gating equations:"
    )
    add_formula_p(
        "f_t = sigmoid( W_f * [h_{t-1}, x_t] + b_f )\n"
        "i_t = sigmoid( W_i * [h_{t-1}, x_t] + b_i )\n"
        "C_tilde_t = tanh( W_c * [h_{t-1}, x_t] + b_c )\n"
        "C_t = f_t * C_{t-1} + i_t * C_tilde_t\n"
        "o_t = sigmoid( W_o * [h_{t-1}, x_t] + b_o )\n"
        "h_t = o_t * tanh( C_t )"
    )
    add_body_p(
        "Where f_t, i_t, o_t represent the forget, input, and output gates respectively, C_t represents the hidden cell state accumulator, and h_t represents the output state."
    )
    
    add_heading_2("3.4 Unified Risk Score Fusion")
    add_body_p(
        "To compile the final risk score for the project, the system performs a weighted linear combination of GNN topological risk predictions, "
        "LSTM temporal anomalies, static analysis outputs, sandboxed run behaviors, and package historical health indexes:"
    )
    add_formula_p("Risk_unified = w_1 * Risk_GNN + w_2 * Risk_LSTM + w_3 * Risk_Static + w_4 * Risk_Behavior + w_5 * Risk_Metadata")
    add_body_p("Where weights are configured such that sum_{i=1}^{5} w_i = 1.0. This guarantees a balanced, robust prediction outcome.")

    # Save
    out_dir = "outputs"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Project_Overview_Threat_Detection_v2.docx")
    doc.save(out_path)
    print(f"[SUCCESS] Word document generated successfully at: {os.path.abspath(out_path)}")

if __name__ == "__main__":
    create_comprehensive_document()
